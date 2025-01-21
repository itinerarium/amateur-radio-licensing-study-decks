#!/usr/bin/env python3

import docx # but pypi package python-docx
import genanki
import html
import re

class AnkiNote(genanki.Note):
    # have question be identified by question ID only for deduplication
    @property
    def guid(self):
        return genanki.guid_for(self.fields[0])

class Question:
    anki_model = genanki.Model(
            1851767301, # arbitrary value
            'BaseModel',
            fields = [
                {'name': 'QuestionID'},
                {'name': 'RegulationReference'},
                {'name': 'Question'},
                {'name': 'Choices'},
                {'name': 'ChoicesWithContrast'},
                {'name': 'Answer'},
                {'name': 'Figure'},
            ],
            templates = [
                {
                    'name': 'ContrastCard',
                    'qfmt': '{{QuestionID}}{{Question}}{{Figure}}{{ChoicesWithContrast}}',
                    'afmt': '{{QuestionID}}{{Question}}{{Figure}}<hr id="answer">{{Answer}}',
                },
            ])

    def __init__(self, q_id: str, regulation: str, question_text: str, choices: list[str], answer_choice_letter: str, answer_text: str, figure: str) -> None:
        self.q_id = q_id
        self.regulation = regulation
        self.question_text = question_text
        self.choices = choices
        self.answer_choice_letter = answer_choice_letter
        self.answer_text = answer_text
        self.figure = figure

    def generate_anki_note(self) -> AnkiNote:
        he_q_id = html.escape(self.q_id)
        he_q_regulation = html.escape(self.regulation)
        he_q_question = html.escape(self.question_text)
        he_q_question = f"<p>{he_q_question}</p>"
        he_q_answer = html.escape(self.answer_text)
        he_q_answer = f"<p><b>{he_q_answer}</b></p>"

        figure = f"""<p><img src="{self.figure}" /></p>""" if self.figure else ''

        choices = ''
        contrast_choices = ''

        j = 0
        while j < len(self.choices):
            he_choice = html.escape(self.choices[j])
            choices += f"<p>{he_choice}</p>"
            if self.choices[j].startswith(self.answer_choice_letter):
                he_choice = f"<b>{he_choice}</b>"
            else:
                he_choice = f"""<span style="opacity: 30%;">{he_choice}</span>"""
            contrast_choices += f"<p>{he_choice}</p>"
            j += 1

        anki_note = AnkiNote(model=self.anki_model, fields=[
                he_q_id,
                he_q_regulation,
                he_q_question,
                choices,
                contrast_choices,
                he_q_answer,
                figure])
        return anki_note

class AnkiDeckMaker:
    ANSWER_LUT = {
        'A': 0,
        'B': 1,
        'C': 2,
        'D': 3
    }

    def extract_question(self, i_paragraph: int) -> Question:
        ref = self.doc.paragraphs[i_paragraph].text
        question = self.doc.paragraphs[i_paragraph+1].text
        choices = [self.doc.paragraphs[i_paragraph+2+j].text for j in range(4)]

        # parse question reference line
        # [prefix] question reference (answer) [regulation reference]
        ref_line = ref.split(' ')
        q_regulation = ''

        if len(ref_line) > 2:
            q_regulation = ' '.join(ref_line[2:])
            q_regulation = q_regulation[1:-1]

        q_id = ref_line[0]

        q_answer_letter = ref_line[1]
        q_answer_letter = q_answer_letter[1:-1]

        q_answer_text = choices[self.ANSWER_LUT[q_answer_letter]][3:]

        q_figure = ''
        if (match := re.search(self.figure_search_re, question)):
            q_figure = self.figure_mapping.get(match.group(1))

        question_entry = Question(
                q_id,
                q_regulation,
                question,
                choices,
                q_answer_letter,
                q_answer_text,
                q_figure)
        return question_entry

    def __init__(self, filepath_input_docx: str, question_prefix: str, figure_mapping: dict[str,str], deck_id: int, deck_title: str, filepath_output_apkg: str) -> None:
        self.filepath_input_docx = filepath_input_docx
        self.question_prefix = question_prefix
        self.figure_mapping = figure_mapping
        self.deck_id = deck_id
        self.deck_title = deck_title
        self.filepath_output_apkg = filepath_output_apkg

        self.figure_search_re = f"[Ff]igure ({self.question_prefix}[A-Z0-9-]+)"

        anki_deck = genanki.Deck(self.deck_id, self.deck_title)

        self.doc = docx.Document(self.filepath_input_docx)
        # this makes certain assumptions about input format and question format (single choice from 4 options)

        # recognize that questions are in blocks of 6 lines
        # starting with question reference, correct answer, regulation reference
        # question
        # choice A
        # choice B
        # choice C
        # choice D

        i = 0
        doc_paragraph_count = len(self.doc.paragraphs)

        while i + 6 < doc_paragraph_count:
            if self.doc.paragraphs[i].text.startswith(self.question_prefix) \
                        and self.doc.paragraphs[i+2].text.startswith('A'):
                question = self.extract_question(i)
                anki_note = question.generate_anki_note()

                anki_deck.add_note(anki_note)
                i += 6
                continue
            i += 1

        anki_deck_package = genanki.Package(anki_deck)
        if self.figure_mapping:
            anki_deck_package.media_files = list(set(self.figure_mapping.values()))
        anki_deck_package.write_to_file(self.filepath_output_apkg)

if __name__ == "__main__":
    technician_figure_mapping = {
            'T-1': 'T1.jpg',
            'T-2': 'T2.jpg',
            'T-3': 'T3.jpg',
            }

    technician_deck_maker = AnkiDeckMaker(
            '20220307-ncvec-technician-exam.docx',
            'T',
            technician_figure_mapping,
            1624516877,
            'Technician Question Pool (6/2026) [eff. 3/2022]',
            '202203-technician-question-pool.apkg',
            )

    general_figure_mapping = {
            'G7-1': 'G7-1.png',
            }

    general_deck_maker = AnkiDeckMaker(
            '20241108-ncvec-general-exam.docx',
            'G',
            general_figure_mapping,
            1495389397,
            'General Question Pool (6/2027) [eff. 11/2024]',
            '202411-general-question-pool.apkg'
            )

    extra_figure_mapping = {
            'E5-1': 'E5-1.png',
            'E6-1': 'E6-1.png',
            'E6-2': 'E6-2.png',
            'E6-3': 'E6-3.png',
            'E7-1': 'E7-1.png',
            'E7-2': 'E7-2.png',
            'E7-3': 'E7-3.png',
            'E9-1': 'E9-1.png',
            'E9-2': 'E9-2.png',
            'E9-3': 'E9-3.png',
            }

    extra_deck_maker = AnkiDeckMaker(
            '20241108-ncvec-extra-exam.docx',
            'E',
            extra_figure_mapping,
            1353326225,
            'Extra Question Pool (2028) [eff. 11/2024]',
            '202411-extra-question-pool.apkg'
            )